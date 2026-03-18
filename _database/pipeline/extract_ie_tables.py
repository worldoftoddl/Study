#!/usr/bin/env python3
"""
K-IFRS 1001 적용사례 테이블 추출 스크립트

DOCX 원본에서 적용사례(실무적용지침) 섹션의 재무제표 예시 테이블을
정확하게 추출하여 마크다운으로 변환합니다.
"""

import io
import zipfile
from pathlib import Path
from docx import Document
from docx.table import Table


def fix_docx_zip(docx_path: str) -> io.BytesIO:
    """DOCX ZIP 내부 경로의 백슬래시를 슬래시로 변환."""
    buf = io.BytesIO()
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                item.filename = item.filename.replace('\\', '/')
                zout.writestr(item, data)
    buf.seek(0)
    return buf


def table_to_md(table: Table, title: str = '') -> str:
    """python-docx Table을 마크다운 테이블로 변환."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ''

    # 빈 구분 열 제거 (모든 행에서 빈 열)
    n_cols = len(rows[0])
    empty_cols = set()
    for col_idx in range(n_cols):
        if all(rows[r][col_idx] == '' for r in range(len(rows))):
            empty_cols.add(col_idx)

    # 열 필터링
    filtered_rows = []
    for row in rows:
        filtered_rows.append([row[i] for i in range(n_cols) if i not in empty_cols])

    if not filtered_rows or not filtered_rows[0]:
        return ''

    # 완전히 빈 행은 제거하되, 구분선으로 빈 행이 필요하면 유지
    clean_rows = []
    for row in filtered_rows:
        if any(cell != '' for cell in row):
            clean_rows.append(row)

    if not clean_rows:
        return ''

    # 마크다운 테이블 생성
    n_cols = len(clean_rows[0])
    lines = []
    if title:
        lines.append(f'\n{title}\n')

    # 헤더 행
    header = clean_rows[0]
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('| ' + ' | '.join(['---'] * n_cols) + ' |')

    # 데이터 행
    for row in clean_rows[1:]:
        # 열 수 맞추기
        while len(row) < n_cols:
            row.append('')
        lines.append('| ' + ' | '.join(row[:n_cols]) + ' |')

    return '\n'.join(lines)


def extract_ie_section(docx_path: str) -> str:
    """적용사례 섹션의 텍스트와 테이블을 추출."""
    buf = fix_docx_zip(docx_path)
    doc = Document(buf)

    # 적용사례 섹션 시작 찾기
    in_ie = False
    ie_started = False
    result_parts = []

    # 문서 요소 순회 (paragraphs + tables 순서 유지)
    body = doc.element.body
    para_idx = 0
    table_idx = 0

    paragraphs = doc.paragraphs
    tables = doc.tables

    # body의 child elements를 순서대로 순회
    from docx.oxml.ns import qn

    table_map = {}
    for i, t in enumerate(tables):
        table_map[id(t._element)] = i

    para_map = {}
    for i, p in enumerate(paragraphs):
        para_map[id(p._element)] = i

    ie_table_start = None  # 적용사례 첫 테이블 인덱스

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            idx = para_map.get(id(child))
            if idx is None:
                continue
            p = paragraphs[idx]
            text = p.text.strip()

            # 적용사례 섹션 시작 감지
            if not in_ie:
                if text == '실무적용지침' and not ie_started:
                    in_ie = True
                    ie_started = True
                    result_parts.append('## 적용사례\n\n실무적용지침\n')
                    continue
                continue

            # 적용사례 종료 감지 (결론도출근거 시작)
            style_name = p.style.name if p.style else ''
            if '결론도출근거' in text and ('Heading' in style_name or text.startswith('결론도출근거')):
                break
            # 또는 BC로 시작하는 문단이 시작되면 (결론도출근거 영역)
            if text.startswith('BC') and text[2:3].isdigit():
                break

            # 텍스트 추가
            if text:
                # 제목 수준 감지
                if 'Heading' in style_name:
                    level = '###'
                    if '2' in style_name:
                        level = '###'
                    elif '3' in style_name:
                        level = '####'
                    result_parts.append(f'\n{level} {text}\n')
                elif text.startswith('제Ⅰ부') or text.startswith('제Ⅱ부') or text.startswith('제Ⅲ부'):
                    result_parts.append(f'\n### {text}\n')
                elif text.startswith('IG'):
                    result_parts.append(f'\n{text}\n')
                elif text.startswith('⑴') or text.startswith('⑵') or text.startswith('⑶') or text.startswith('⑷'):
                    result_parts.append(f'\n{text}\n')
                else:
                    result_parts.append(f'\n{text}\n')

        elif tag == 'tbl':
            if not in_ie:
                continue
            tidx = table_map.get(id(child))
            if tidx is None:
                continue
            t = tables[tidx]
            md = table_to_md(t)
            if md:
                result_parts.append(f'\n{md}\n')

    return '\n'.join(result_parts)


def main():
    docx_dir = Path(__file__).parent.parent / 'data' / 'raw' / 'IFRS_docx'
    docx_file = list(docx_dir.glob('*제1001호*'))[0]

    print(f"원본 DOCX: {docx_file.name}")
    content = extract_ie_section(str(docx_file))

    out_dir = Path(__file__).parent.parent / 'output' / 'docx_md' / '적용사례'
    out_dir.mkdir(exist_ok=True)
    out_file = out_dir / 'IE_K-IFRS_1001_재무제표_표시.md'
    out_file.write_text(content, encoding='utf-8')

    line_count = content.count('\n')
    print(f"출력: {out_file}")
    print(f"총 {line_count}줄")


if __name__ == '__main__':
    main()
